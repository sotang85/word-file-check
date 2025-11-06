"""Tests ensuring tables are processed when extracting sentences and building docs."""
from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

try:
    from docx import Document  # type: ignore

    DOCX_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover - exercised in environments without python-docx
    DOCX_AVAILABLE = False

from lexdiff import build_highlighted_document, compare_sentences, load_sentences


@unittest.skipUnless(DOCX_AVAILABLE, "requires python-docx")
class TableSupportTests(unittest.TestCase):
    def test_load_sentences_includes_table_cells(self) -> None:
        with TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "table.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "표 문장 첫째. 표 문장 둘째."
            document.save(path)

            sentences = load_sentences(str(path))
            texts = [sentence.text for sentence in sentences]

            self.assertTrue(any("표 문장 첫째" in text for text in texts))
            self.assertTrue(any("표 문장 둘째" in text for text in texts))

            table_sentence = next(sentence for sentence in sentences if sentence.text == "표 문장 둘째")
            self.assertEqual(table_sentence.container, "table")
            self.assertIsNotNone(table_sentence.table_index)
            self.assertIsNotNone(table_sentence.row_index)
            self.assertIsNotNone(table_sentence.cell_index)

    def test_highlight_document_creates_table_output(self) -> None:
        with TemporaryDirectory() as tmpdir:
            source_path = Path(tmpdir) / "source.docx"
            target_path = Path(tmpdir) / "target.docx"
            output_path = Path(tmpdir) / "result.docx"

            source = Document()
            src_table = source.add_table(rows=1, cols=1)
            src_table.cell(0, 0).text = "첫째 문장. 둘째 문장."
            source.save(source_path)

            target = Document()
            tgt_table = target.add_table(rows=1, cols=1)
            tgt_table.cell(0, 0).text = "첫째 문장. 둘째 문장 변경."
            target.save(target_path)

            sentences_a = load_sentences(str(source_path))
            sentences_b = load_sentences(str(target_path))

            operations = compare_sentences(sentences_a, sentences_b, ignore=[], threshold=0.8)
            build_highlighted_document(operations, str(output_path))

            result = Document(output_path)
            self.assertGreaterEqual(len(result.tables), 1)
            combined_text = " ".join(cell.text for row in result.tables[0].rows for cell in row.cells)
            self.assertIn("둘째 문장", combined_text)


if __name__ == "__main__":  # pragma: no cover
    unittest.main()
