#!/usr/bin/env python3
"""CLI tool to highlight differences between two DOCX files."""
from __future__ import annotations

import argparse
import csv
import os
import re
import string
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from difflib import SequenceMatcher
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


@dataclass
class SentenceRecord:
    index: int
    text: str
    paragraph_index: int
    sentence_in_paragraph: int
    prefix: str = ""
    postfix: str = ""


EXTRA_PUNCTUATION = "“”‘’‚‛„‟‹›«»、，；：·…‧〈〉《》「」『』【】〔〕（）［］｛｝()[]{}<>？！。．﹒﹔﹖﹗"
PUNCTUATION_TRANSLATION = str.maketrans('', '', string.punctuation + EXTRA_PUNCTUATION)


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Compare DOCX files at sentence level.")
    parser.add_argument("source", help="Original DOCX document")
    parser.add_argument("target", help="Revised DOCX document")
    parser.add_argument("--out", dest="out_docx", required=True, help="Path to highlighted DOCX output")
    parser.add_argument("--csv", dest="out_csv", required=True, help="Path to CSV diff report")
    parser.add_argument(
        "--ignore",
        default="",
        help="Comma separated list of elements to ignore when comparing (options: punct, space)",
    )
    parser.add_argument(
        "--threshold",
        type=float,
        default=0.8,
        help="Similarity threshold (0-1) for classifying replacements",
    )
    args = parser.parse_args(argv)

    args.ignore_tokens = [token.strip().lower() for token in args.ignore.split(",") if token.strip()]
    invalid_tokens = sorted(set(args.ignore_tokens) - {"punct", "space"})
    if invalid_tokens:
        parser.error(f"Unsupported ignore options: {', '.join(invalid_tokens)}")
    if not 0 <= args.threshold <= 1:
        parser.error("--threshold must be between 0 and 1")

    return args


def load_sentences(path: str) -> List[SentenceRecord]:
    document = Document(path)
    sentences: List[SentenceRecord] = []
    idx = 0
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        raw_text = paragraph.text.replace("\r", "")
        if not raw_text.strip():
            continue
        segments = split_paragraph_segments(raw_text)
        for sentence_idx, (prefix, content, postfix) in enumerate(segments):
            trimmed = content.strip()
            if not trimmed:
                continue
            sentences.append(
                SentenceRecord(
                    index=idx,
                    text=trimmed,
                    paragraph_index=paragraph_index,
                    sentence_in_paragraph=sentence_idx,
                    prefix=prefix,
                    postfix=postfix,
                )
            )
            idx += 1
    return sentences


SENTENCE_PATTERN = re.compile(r"[^\n.!?。！？]+(?:[.!?。！？]+(?=\s|$)|$)")


def split_paragraph_segments(text: str) -> List[Tuple[str, str, str]]:
    cleaned = text.replace("\r", "")
    matches = list(SENTENCE_PATTERN.finditer(cleaned))
    segments: List[Tuple[str, str, str]] = []
    cursor = 0

    if not matches:
        stripped = cleaned.strip()
        if stripped:
            leading = cleaned[: len(cleaned) - len(cleaned.lstrip())]
            trailing = cleaned[len(cleaned.rstrip()):]
            segments.append((leading, stripped, trailing))
        return segments

    for match in matches:
        start, end = match.span()
        prefix = cleaned[cursor:start]
        content = match.group()
        leading_inner_len = len(content) - len(content.lstrip())
        trailing_inner_len = len(content.rstrip()) - len(content.strip())
        leading_inner = content[:leading_inner_len]
        trailing_inner = content[len(content) - trailing_inner_len :] if trailing_inner_len else ""
        core_start = leading_inner_len
        core_end = len(content) - trailing_inner_len
        core = content[core_start:core_end] if core_end > core_start else ""
        follow_ws = re.match(r"\s*", cleaned[end:]).group()
        cursor = end + len(follow_ws)
        core_text = core.strip()
        segments.append((prefix + leading_inner, core_text, trailing_inner + follow_ws))

    if cursor < len(cleaned) and segments:
        last_prefix, last_core, last_postfix = segments[-1]
        segments[-1] = (last_prefix, last_core, last_postfix + cleaned[cursor:])

    return segments


def normalize_text(text: str, ignore: Iterable[str]) -> str:
    normalized = text
    if "space" in ignore:
        normalized = re.sub(r"\s+", "", normalized)
    if "punct" in ignore:
        normalized = normalized.translate(PUNCTUATION_TRANSLATION)
    return normalized


def extract_operations(
    sentences_a: List[SentenceRecord],
    sentences_b: List[SentenceRecord],
    ignore: Iterable[str],
    threshold: float,
) -> List[dict]:
    ignore_set = set(x.strip().lower() for x in ignore if x.strip())
    norm_a = [normalize_text(rec.text, ignore_set) for rec in sentences_a]
    norm_b = [normalize_text(rec.text, ignore_set) for rec in sentences_b]
    matcher = SequenceMatcher(None, norm_a, norm_b, autojunk=False)
    operations: List[dict] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i1, i2):
                operations.append(
                    {
                        "type": "equal",
                        "a": sentences_a[offset],
                        "b": sentences_b[j1 + (offset - i1)],
                        "sim": 1.0,
                    }
                )
        elif tag == "delete":
            for offset in range(i1, i2):
                operations.append({"type": "del", "a": sentences_a[offset], "b": None, "sim": 0.0})
        elif tag == "insert":
            for offset in range(j1, j2):
                operations.append({"type": "add", "a": None, "b": sentences_b[offset], "sim": 0.0})
        elif tag == "replace":
            segment_a = list(range(i1, i2))
            segment_b = list(range(j1, j2))
            length = min(len(segment_a), len(segment_b))
            for idx in range(length):
                rec_a = sentences_a[segment_a[idx]]
                rec_b = sentences_b[segment_b[idx]]
                norm_a_sentence = normalize_text(rec_a.text, ignore_set)
                norm_b_sentence = normalize_text(rec_b.text, ignore_set)
                if norm_a_sentence == norm_b_sentence:
                    operations.append({"type": "equal", "a": rec_a, "b": rec_b, "sim": 1.0})
                    continue
                ratio = SequenceMatcher(None, norm_a_sentence, norm_b_sentence, autojunk=False).ratio()
                if ratio >= threshold:
                    operations.append({"type": "replace", "a": rec_a, "b": rec_b, "sim": ratio})
                else:
                    operations.append({"type": "del", "a": rec_a, "b": None, "sim": 0.0})
                    operations.append({"type": "add", "a": None, "b": rec_b, "sim": 0.0})
            if len(segment_a) > length:
                for idx in segment_a[length:]:
                    operations.append({"type": "del", "a": sentences_a[idx], "b": None, "sim": 0.0})
            if len(segment_b) > length:
                for idx in segment_b[length:]:
                    operations.append({"type": "add", "a": None, "b": sentences_b[idx], "sim": 0.0})
    return operations


def annotate_numeric_delta(original: str, revised: str) -> str:
    number_pattern = re.compile(r"-?\d+(?:\.\d+)?")
    numbers_original = number_pattern.findall(original)
    numbers_revised = number_pattern.findall(revised)

    if not numbers_original and not numbers_revised:
        return revised

    def to_decimal(value: str) -> Optional[Decimal]:
        try:
            return Decimal(value)
        except (InvalidOperation, ValueError):
            return None

    def format_delta(delta: Decimal) -> str:
        if delta == 0:
            return "0"
        prefix = "+" if delta > 0 else ""
        normalized = delta.normalize()
        if normalized == normalized.to_integral():
            return f"{prefix}{int(normalized)}"
        return f"{prefix}{str(normalized)}"

    deltas: List[str] = []
    changed = False

    paired = zip(numbers_original, numbers_revised)
    for original_value, revised_value in paired:
        original_decimal = to_decimal(original_value)
        revised_decimal = to_decimal(revised_value)
        if original_decimal is None or revised_decimal is None:
            continue
        if original_decimal != revised_decimal:
            changed = True
        delta = revised_decimal - original_decimal
        deltas.append(format_delta(delta))

    if len(numbers_original) > len(numbers_revised):
        changed = True
        for removed in numbers_original[len(numbers_revised):]:
            deltas.append(f"-{removed} (removed)")
    elif len(numbers_revised) > len(numbers_original):
        changed = True
        for added in numbers_revised[len(numbers_original):]:
            deltas.append(f"+{added} (new)")

    if changed and deltas:
        return f"{revised} (Δ {', '.join(deltas)})"
    return revised


def tokenize(text: str) -> List[str]:
    token_pattern = re.compile(r"\s+|[\w\-\u00C0-\u02AF\u0400-\u04FF\uAC00-\uD7AF]+|[^\w\s]", re.UNICODE)
    tokens = token_pattern.findall(text)
    if not tokens:
        return [text]
    return tokens


def append_text(paragraph, text: str, formatter: Optional[Callable[[Any], None]] = None) -> None:
    if not text:
        return
    for segment in re.split(r"(\n)", text):
        if not segment:
            continue
        if segment == "\n":
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(segment)
        if formatter:
            formatter(run)


def build_highlighted_document(operations: List[dict], output_path: str) -> None:
    document = Document()
    paragraph_cache: Dict[int, Any] = {}
    highest_created = -1

    def ensure_paragraph(paragraph_index: int):
        nonlocal highest_created
        if paragraph_index in paragraph_cache:
            return paragraph_cache[paragraph_index]
        while highest_created < paragraph_index:
            highest_created += 1
            paragraph_cache[highest_created] = document.add_paragraph()
        return paragraph_cache[paragraph_index]

    for op in operations:
        op_type = op["type"]
        if op_type == "del":
            paragraph = document.add_paragraph()
            record_a = op.get("a")
            if record_a:
                if record_a.prefix:
                    append_text(paragraph, record_a.prefix)
                append_text(paragraph, record_a.text, lambda run: setattr(run.font, "strike", True))
                if record_a.postfix:
                    append_text(paragraph, record_a.postfix)
            if not paragraph.text:
                paragraph.add_run("\u00A0")
            continue

        record = op.get("b") or op.get("a")
        if record is None:
            continue
        paragraph = ensure_paragraph(record.paragraph_index)

        if record.prefix:
            append_text(paragraph, record.prefix)

        if op_type == "add":
            append_text(paragraph, record.text, lambda run: setattr(run.font, "underline", True))
        elif op_type == "equal":
            append_text(paragraph, record.text)
        elif op_type == "replace":
            rec_a = op.get("a")
            rec_b = op.get("b")
            tokens_a = tokenize(rec_a.text if rec_a else "")
            tokens_b = tokenize(rec_b.text if rec_b else "")
            matcher = SequenceMatcher(None, tokens_a, tokens_b, autojunk=False)
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                segment = "".join(tokens_b[j1:j2])
                if not segment:
                    continue
                if tag in {"replace", "insert"}:
                    append_text(paragraph, segment, lambda run: setattr(run.font, "highlight_color", WD_COLOR_INDEX.YELLOW))
                else:
                    append_text(paragraph, segment)
        if record.postfix:
            append_text(paragraph, record.postfix)

    document.save(output_path)


def write_csv_report(operations: List[dict], output_path: str) -> None:
    fieldnames = ["type", "sim", "original", "revised", "idxA", "idxB"]
    with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for op in operations:
            if op["type"] == "equal":
                continue
            original = op["a"].text if op.get("a") else ""
            revised = op["b"].text if op.get("b") else ""
            if op["type"] == "replace":
                revised = annotate_numeric_delta(original, revised)
            writer.writerow(
                {
                    "type": op["type"],
                    "sim": f"{op.get('sim', 0.0):.2f}",
                    "original": original,
                    "revised": revised,
                    "idxA": op["a"].index + 1 if op.get("a") else "",
                    "idxB": op["b"].index + 1 if op.get("b") else "",
                }
            )


def run_diff(
    source: str,
    target: str,
    out_docx: str,
    out_csv: str,
    ignore_tokens: Optional[Iterable[str]] = None,
    threshold: float = 0.8,
) -> List[dict]:
    """Execute the diff workflow without relying on CLI parsing.

    Returns the list of diff operations so callers such as the GUI can present
    the result immediately without reparsing the documents.
    """

    ignore_list = [token.strip().lower() for token in (ignore_tokens or []) if token.strip()]

    if not os.path.exists(source):
        raise FileNotFoundError(f"Source file not found: {source}")
    if not os.path.exists(target):
        raise FileNotFoundError(f"Target file not found: {target}")

    sentences_a = load_sentences(source)
    sentences_b = load_sentences(target)

    operations = extract_operations(sentences_a, sentences_b, ignore_list, threshold)

    for path in (out_docx, out_csv):
        directory = os.path.dirname(os.path.abspath(path))
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

    build_highlighted_document(operations, out_docx)
    write_csv_report(operations, out_csv)

    return operations


def main(argv: Optional[Sequence[str]] = None) -> None:
    args = parse_args(argv)

    run_diff(
        source=args.source,
        target=args.target,
        out_docx=args.out_docx,
        out_csv=args.out_csv,
        ignore_tokens=args.ignore_tokens,
        threshold=args.threshold,
    )


if __name__ == "__main__":
    main()
