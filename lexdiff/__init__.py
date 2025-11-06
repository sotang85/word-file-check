"""Core diff engine shared by the CLI, GUI, and web UI."""
from __future__ import annotations

import csv
import os
import re
import string
from dataclasses import dataclass, field
from decimal import Decimal, InvalidOperation
from difflib import SequenceMatcher
from typing import Iterable, List, Optional, Sequence, Tuple, Dict, Literal

__all__ = [
    "DependencyError",
    "Sentence",
    "WordDiff",
    "Operation",
    "DiffRow",
    "DiffResult",
    "parse_ignore_tokens",
    "load_sentences",
    "compare_sentences",
    "build_highlighted_document",
    "build_csv_rows",
    "write_csv",
    "run_diff",
    "annotate_numeric_delta",
]


class DependencyError(RuntimeError):
    """Raised when a required optional dependency is missing."""


@dataclass
class Sentence:
    """Sentence metadata captured from a DOCX paragraph."""

    index: int
    text: str
    paragraph_index: int
    sentence_in_paragraph: int
    prefix: str = ""
    postfix: str = ""


@dataclass
class WordDiff:
    """Token-level change description for replacement operations."""

    kind: Literal["equal", "insert", "delete", "replace"]
    original: str
    revised: str


@dataclass
class Operation:
    """Single diff operation at the sentence level."""

    kind: Literal["equal", "add", "del", "replace"]
    similarity: float
    original: Optional[Sentence] = None
    revised: Optional[Sentence] = None
    word_diff: List[WordDiff] = field(default_factory=list)


@dataclass
class DiffRow:
    """CSV representation of a change row."""

    type: str
    sim: str
    original: str
    revised: str
    idxA: str
    idxB: str

    def to_dict(self) -> Dict[str, str]:
        return {
            "type": self.type,
            "sim": self.sim,
            "original": self.original,
            "revised": self.revised,
            "idxA": self.idxA,
            "idxB": self.idxB,
        }


@dataclass
class DiffResult:
    """Aggregate object returned by :func:`run_diff`."""

    operations: List[Operation]
    rows: List[DiffRow]


EXTRA_PUNCTUATION = "“”‘’‚‛„‟‹›«»、，；：·…‧〈〉《》「」『』【】〔〕（）［］｛｝()[]{}<>？！。．﹒﹔﹖﹗"
PUNCTUATION_TRANSLATION = str.maketrans("", "", string.punctuation + EXTRA_PUNCTUATION)

SENTENCE_PATTERN = re.compile(r"[^\n.!?。！？]+(?:[.!?。！？]+(?=\s|$)|$)")
TOKEN_PATTERN = re.compile(r"\s+|[\w\-\u00C0-\u02AF\u0400-\u04FF\uAC00-\uD7AF]+|[^\w\s]", re.UNICODE)
NUMBER_PATTERN = re.compile(r"-?\d[\d,]*(?:\.\d+)?")


def _require_docx():  # type: ignore[override]
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_COLOR_INDEX  # type: ignore
    except ModuleNotFoundError as exc:  # pragma: no cover - exercised via runtime error handling
        raise DependencyError(
            "python-docx 패키지가 설치되어 있지 않습니다. `pip install python-docx` 명령으로 설치해 주세요."
        ) from exc
    return Document, WD_COLOR_INDEX


def parse_ignore_tokens(values: Iterable[str]) -> List[str]:
    tokens = [token.strip().lower() for token in values if token and token.strip()]
    invalid = sorted(set(tokens) - {"punct", "space"})
    if invalid:
        raise ValueError(f"지원하지 않는 ignore 옵션입니다: {', '.join(invalid)}")
    return tokens


def _split_paragraph_segments(text: str) -> List[Tuple[str, str, str]]:
    cleaned = text.replace("\r", "")
    matches = list(SENTENCE_PATTERN.finditer(cleaned))
    segments: List[Tuple[str, str, str]] = []
    cursor = 0

    if not matches:
        stripped = cleaned.strip()
        if stripped:
            leading = cleaned[: len(cleaned) - len(cleaned.lstrip())]
            trailing = cleaned[len(cleaned.rstrip()):]
            segments.append((leading, stripped, trailing))
        return segments

    for match in matches:
        start, end = match.span()
        prefix = cleaned[cursor:start]
        content = match.group()
        leading_inner_len = len(content) - len(content.lstrip())
        trailing_inner_len = len(content.rstrip()) - len(content.strip())
        leading_inner = content[:leading_inner_len]
        trailing_inner = content[len(content) - trailing_inner_len :] if trailing_inner_len else ""
        core_start = leading_inner_len
        core_end = len(content) - trailing_inner_len
        core = content[core_start:core_end] if core_end > core_start else ""
        follow_ws = re.match(r"\s*", cleaned[end:]).group()
        cursor = end + len(follow_ws)
        core_text = core.strip()
        segments.append((prefix + leading_inner, core_text, trailing_inner + follow_ws))

    if cursor < len(cleaned) and segments:
        last_prefix, last_core, last_postfix = segments[-1]
        segments[-1] = (last_prefix, last_core, last_postfix + cleaned[cursor:])

    return segments


def load_sentences(path: str) -> List[Sentence]:
    Document, _ = _require_docx()
    document = Document(path)
    sentences: List[Sentence] = []
    idx = 0
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        raw_text = paragraph.text.replace("\r", "")
        if not raw_text.strip():
            continue
        segments = _split_paragraph_segments(raw_text)
        for sentence_idx, (prefix, content, postfix) in enumerate(segments):
            trimmed = content.strip()
            if not trimmed:
                continue
            sentences.append(
                Sentence(
                    index=idx,
                    text=trimmed,
                    paragraph_index=paragraph_index,
                    sentence_in_paragraph=sentence_idx,
                    prefix=prefix,
                    postfix=postfix,
                )
            )
            idx += 1
    return sentences


def _normalize_text(text: str, ignore: Iterable[str]) -> str:
    normalized = text
    if "space" in ignore:
        normalized = re.sub(r"\s+", "", normalized)
    if "punct" in ignore:
        normalized = normalized.translate(PUNCTUATION_TRANSLATION)
    return normalized


def _tokenize(text: str) -> List[str]:
    tokens = TOKEN_PATTERN.findall(text)
    return tokens or [text]


def _build_word_diff(original: str, revised: str) -> List[WordDiff]:
    tokens_a = _tokenize(original)
    tokens_b = _tokenize(revised)
    matcher = SequenceMatcher(None, tokens_a, tokens_b, autojunk=False)
    segments: List[WordDiff] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        text_a = "".join(tokens_a[i1:i2])
        text_b = "".join(tokens_b[j1:j2])
        if tag == "equal" and text_b:
            segments.append(WordDiff("equal", text_a, text_b))
        elif tag == "insert" and text_b:
            segments.append(WordDiff("insert", text_a, text_b))
        elif tag == "delete" and text_a:
            segments.append(WordDiff("delete", text_a, text_b))
        elif tag == "replace" and (text_a or text_b):
            segments.append(WordDiff("replace", text_a, text_b))
    return segments


def compare_sentences(
    sentences_a: List[Sentence],
    sentences_b: List[Sentence],
    ignore: Iterable[str],
    threshold: float,
) -> List[Operation]:
    ignore_set = set(ignore)
    norm_a = [_normalize_text(rec.text, ignore_set) for rec in sentences_a]
    norm_b = [_normalize_text(rec.text, ignore_set) for rec in sentences_b]
    matcher = SequenceMatcher(None, norm_a, norm_b, autojunk=False)
    operations: List[Operation] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i1, i2):
                operations.append(
                    Operation(
                        kind="equal",
                        similarity=1.0,
                        original=sentences_a[offset],
                        revised=sentences_b[j1 + (offset - i1)],
                    )
                )
        elif tag == "delete":
            for offset in range(i1, i2):
                operations.append(Operation(kind="del", similarity=0.0, original=sentences_a[offset]))
        elif tag == "insert":
            for offset in range(j1, j2):
                operations.append(Operation(kind="add", similarity=0.0, revised=sentences_b[offset]))
        elif tag == "replace":
            segment_a = list(range(i1, i2))
            segment_b = list(range(j1, j2))
            length = min(len(segment_a), len(segment_b))
            for idx in range(length):
                rec_a = sentences_a[segment_a[idx]]
                rec_b = sentences_b[segment_b[idx]]
                norm_a_sentence = _normalize_text(rec_a.text, ignore_set)
                norm_b_sentence = _normalize_text(rec_b.text, ignore_set)
                if norm_a_sentence == norm_b_sentence:
                    operations.append(Operation(kind="equal", similarity=1.0, original=rec_a, revised=rec_b))
                    continue
                ratio = SequenceMatcher(None, norm_a_sentence, norm_b_sentence, autojunk=False).ratio()
                if ratio >= threshold:
                    word_diff = _build_word_diff(rec_a.text, rec_b.text)
                    operations.append(
                        Operation(
                            kind="replace",
                            similarity=ratio,
                            original=rec_a,
                            revised=rec_b,
                            word_diff=word_diff,
                        )
                    )
                else:
                    operations.append(Operation(kind="del", similarity=0.0, original=rec_a))
                    operations.append(Operation(kind="add", similarity=0.0, revised=rec_b))
            if len(segment_a) > length:
                for idx in segment_a[length:]:
                    operations.append(Operation(kind="del", similarity=0.0, original=sentences_a[idx]))
            if len(segment_b) > length:
                for idx in segment_b[length:]:
                    operations.append(Operation(kind="add", similarity=0.0, revised=sentences_b[idx]))
    return operations


def _ensure_directory(path: str) -> None:
    directory = os.path.dirname(os.path.abspath(path))
    if directory and not os.path.exists(directory):
        os.makedirs(directory, exist_ok=True)


def _apply_formatting(run, **flags) -> None:
    for key, value in flags.items():
        setattr(run.font, key, value)


def _append_text(paragraph, text: str, formatter=None) -> None:
    if not text:
        return
    for segment in re.split(r"(\n)", text):
        if not segment:
            continue
        if segment == "\n":
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(segment)
        if formatter:
            formatter(run)


def build_highlighted_document(operations: Sequence[Operation], output_path: str) -> None:
    Document, WD_COLOR_INDEX = _require_docx()
    document = Document()
    paragraph_cache: Dict[int, object] = {}
    highest_created = -1

    def ensure_paragraph(paragraph_index: int):
        nonlocal highest_created
        if paragraph_index in paragraph_cache:
            return paragraph_cache[paragraph_index]
        while highest_created < paragraph_index:
            highest_created += 1
            paragraph_cache[highest_created] = document.add_paragraph()
        return paragraph_cache[paragraph_index]

    for op in operations:
        if op.kind == "del":
            paragraph = document.add_paragraph()
            record = op.original
            if record:
                if record.prefix:
                    _append_text(paragraph, record.prefix)
                _append_text(paragraph, record.text, lambda run: _apply_formatting(run, strike=True))
                if record.postfix:
                    _append_text(paragraph, record.postfix)
            if not paragraph.text:
                paragraph.add_run("\u00A0")
            continue

        record = op.revised or op.original
        if record is None:
            continue
        paragraph = ensure_paragraph(record.paragraph_index)

        if record.prefix:
            _append_text(paragraph, record.prefix)

        if op.kind == "add":
            _append_text(paragraph, record.text, lambda run: _apply_formatting(run, underline=True))
        elif op.kind == "equal":
            _append_text(paragraph, record.text)
        elif op.kind == "replace":
            segments = op.word_diff or _build_word_diff(op.original.text if op.original else "", record.text)
            for segment in segments:
                if not segment.revised:
                    continue
                if segment.kind in {"insert", "replace"}:
                    _append_text(
                        paragraph,
                        segment.revised,
                        lambda run: _apply_formatting(run, highlight_color=WD_COLOR_INDEX.YELLOW),
                    )
                else:
                    _append_text(paragraph, segment.revised)
        else:
            _append_text(paragraph, record.text)

        if record.postfix:
            _append_text(paragraph, record.postfix)

    document.save(output_path)


def annotate_numeric_delta(original: str, revised: str) -> str:
    numbers_original = NUMBER_PATTERN.findall(original)
    numbers_revised = NUMBER_PATTERN.findall(revised)

    if not numbers_original and not numbers_revised:
        return revised

    def to_decimal(value: str) -> Optional[Decimal]:
        try:
            cleaned = value.replace(",", "")
            return Decimal(cleaned)
        except (InvalidOperation, ValueError):
            return None

    def format_delta(delta: Decimal) -> str:
        if delta == 0:
            return "0"
        prefix = "+" if delta > 0 else ""
        normalized = delta.normalize()
        if normalized == normalized.to_integral():
            return f"{prefix}{int(normalized)}"
        return f"{prefix}{str(normalized)}"

    deltas: List[str] = []
    changed = False

    for original_value, revised_value in zip(numbers_original, numbers_revised):
        original_decimal = to_decimal(original_value)
        revised_decimal = to_decimal(revised_value)
        if original_decimal is None or revised_decimal is None:
            continue
        if original_decimal != revised_decimal:
            changed = True
        delta = revised_decimal - original_decimal
        deltas.append(format_delta(delta))

    if len(numbers_original) > len(numbers_revised):
        changed = True
        for removed in numbers_original[len(numbers_revised):]:
            deltas.append(f"-{removed} (removed)")
    elif len(numbers_revised) > len(numbers_original):
        changed = True
        for added in numbers_revised[len(numbers_original):]:
            deltas.append(f"+{added} (new)")

    if changed and deltas:
        return f"{revised} (Δ {', '.join(deltas)})"
    return revised


def build_csv_rows(operations: Iterable[Operation]) -> List[DiffRow]:
    rows: List[DiffRow] = []
    for op in operations:
        if op.kind == "equal":
            continue
        original_text = op.original.text if op.original else ""
        revised_text = op.revised.text if op.revised else ""
        if op.kind == "replace":
            revised_text = annotate_numeric_delta(original_text, revised_text)
        rows.append(
            DiffRow(
                type=op.kind,
                sim=f"{op.similarity:.2f}",
                original=original_text,
                revised=revised_text,
                idxA=str(op.original.index + 1) if op.original else "",
                idxB=str(op.revised.index + 1) if op.revised else "",
            )
        )
    return rows


def write_csv(rows: Sequence[DiffRow], output_path: str) -> None:
    fieldnames = ["type", "sim", "original", "revised", "idxA", "idxB"]
    with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(row.to_dict())


def run_diff(
    source: str,
    target: str,
    out_docx: str,
    out_csv: str,
    ignore_tokens: Optional[Iterable[str]] = None,
    threshold: float = 0.8,
) -> DiffResult:
    ignore_list = parse_ignore_tokens(ignore_tokens or [])

    if not os.path.exists(source):
        raise FileNotFoundError(f"원본 파일을 찾을 수 없습니다: {source}")
    if not os.path.exists(target):
        raise FileNotFoundError(f"수정 파일을 찾을 수 없습니다: {target}")

    sentences_a = load_sentences(source)
    sentences_b = load_sentences(target)

    operations = compare_sentences(sentences_a, sentences_b, ignore_list, threshold)
    rows = build_csv_rows(operations)

    for path in (out_docx, out_csv):
        _ensure_directory(path)

    build_highlighted_document(operations, out_docx)
    write_csv(rows, out_csv)

    return DiffResult(operations=operations, rows=rows)
